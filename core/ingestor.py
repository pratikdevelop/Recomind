"""
core/ingestor.py — Multi-Format Document Ingestor
Supports: .txt  .md  .pdf  .docx  .csv  .json  .html
"""

import csv
import json
import logging
from io import BytesIO, StringIO
from pathlib import Path
from typing import List, Dict, Any, Optional

log = logging.getLogger(__name__)

# ── Format handlers ──────────────────────────────────────────────────────────

def _read_txt(data: bytes, filename: str) -> str:
    return data.decode("utf-8", errors="replace")


def _read_md(data: bytes, filename: str) -> str:
    return data.decode("utf-8", errors="replace")


def _read_pdf(data: bytes, filename: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(data))
        pages  = [p.extract_text() or "" for p in reader.pages]
        return "\n\n".join(pages)
    except ImportError:
        log.error("pypdf not installed — run: pip install pypdf")
        return ""
    except Exception as exc:
        log.error(f"PDF read error ({filename}): {exc}")
        return ""


def _read_docx(data: bytes, filename: str) -> str:
    try:
        from docx import Document
        doc  = Document(BytesIO(data))
        # Paragraphs + table cells
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts)
    except ImportError:
        log.error("python-docx not installed — run: pip install python-docx")
        return ""
    except Exception as exc:
        log.error(f"DOCX read error ({filename}): {exc}")
        return ""


def _read_csv(data: bytes, filename: str) -> str:
    try:
        text    = data.decode("utf-8", errors="replace")
        reader  = csv.DictReader(StringIO(text))
        rows    = list(reader)
        if not rows:
            return ""
        headers = list(rows[0].keys())
        lines   = ["Columns: " + ", ".join(headers)]
        for row in rows:
            lines.append(" | ".join(f"{k}: {v}" for k, v in row.items() if v))
        return "\n".join(lines)
    except Exception as exc:
        log.error(f"CSV read error ({filename}): {exc}")
        return ""


def _read_json(data: bytes, filename: str) -> str:
    try:
        obj = json.loads(data.decode("utf-8", errors="replace"))
        # Flatten to a readable string
        if isinstance(obj, list):
            lines = []
            for item in obj:
                if isinstance(item, dict):
                    lines.append(", ".join(f"{k}: {v}" for k, v in item.items()))
                else:
                    lines.append(str(item))
            return "\n".join(lines)
        return json.dumps(obj, indent=2)
    except Exception as exc:
        log.error(f"JSON read error ({filename}): {exc}")
        return ""


def _read_html(data: bytes, filename: str) -> str:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(data, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "head"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        # Fallback: strip tags naively
        import re
        text = data.decode("utf-8", errors="replace")
        return re.sub(r"<[^>]+>", "", text)
    except Exception as exc:
        log.error(f"HTML read error ({filename}): {exc}")
        return ""


# ── Extension map ─────────────────────────────────────────────────────────────
HANDLERS = {
    ".txt":  _read_txt,
    ".md":   _read_md,
    ".pdf":  _read_pdf,
    ".docx": _read_docx,
    ".csv":  _read_csv,
    ".json": _read_json,
    ".html": _read_html,
    ".htm":  _read_html,
}

SUPPORTED_EXTENSIONS = list(HANDLERS.keys())


def extract_text(data: bytes, filename: str) -> Optional[str]:
    """
    Extract plain text from file bytes.

    Parameters
    ----------
    data     : raw file bytes
    filename : original filename (used to detect format)

    Returns
    -------
    Extracted text string or None if unsupported.
    """
    ext     = Path(filename).suffix.lower()
    handler = HANDLERS.get(ext)
    if not handler:
        log.warning(f"Unsupported file type: {ext} ({filename})")
        return None
    text = handler(data, filename)
    return text.strip() if text else None


def file_to_document(data: bytes, filename: str,
                      extra_metadata: Optional[Dict] = None) -> Optional[Dict[str, Any]]:
    """
    Convert raw file bytes into a document dict ready for insert_documents().
    """
    text = extract_text(data, filename)
    if not text:
        return None

    ext = Path(filename).suffix.lstrip(".")
    metadata = {"source": filename, "type": ext}
    if extra_metadata:
        metadata.update(extra_metadata)

    return {"content": text, "metadata": metadata}


def load_folder(folder: str = "documents") -> List[Dict[str, Any]]:
    """
    Walk a folder and return all extractable documents.
    """
    root = Path(folder)
    if not root.is_dir():
        log.warning(f"Folder '{folder}' not found.")
        return []

    docs  = []
    files = [f for f in root.rglob("*") if f.suffix.lower() in HANDLERS]
    log.info(f"Found {len(files)} file(s) in '{folder}'")

    for fp in files:
        try:
            doc = file_to_document(fp.read_bytes(), fp.name,
                                   extra_metadata={"path": str(fp)})
            if doc:
                docs.append(doc)
                log.info(f"  ✓ {fp.name}")
            else:
                log.warning(f"  ✗ Empty: {fp.name}")
        except Exception as exc:
            log.warning(f"  ✗ Error reading {fp.name}: {exc}")

    return docs